# -*- coding: utf-8 -*-
"""
Markdown 转 Word 脚本（专用于条码打印系统解决方案）
支持：标题 H1-H4、段落、加粗、无序列表、有序列表、表格、代码块、引用块、水平线
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------- 中文字体辅助 ----------------
def set_run_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_runs_with_bold(paragraph, text, base_size=10.5, base_color=None, bold_color=None):
    """把含 **加粗** 的文本，拆成多个 run 写入段落。"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True, color=bold_color or base_color)
        else:
            # 处理行内代码 `code`
            sub_parts = re.split(r'(`[^`]+`)', part)
            for sp in sub_parts:
                if not sp:
                    continue
                if sp.startswith('`') and sp.endswith('`'):
                    run = paragraph.add_run(sp[1:-1])
                    set_run_font(run, name='Consolas', size=base_size - 0.5)
                    # 给行内代码加灰色背景
                    rPr = run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    rPr.append(shd)
                else:
                    run = paragraph.add_run(sp)
                    set_run_font(run, size=base_size, color=base_color)


# ---------------- 解析与写入 ----------------
def convert(md_path: str, docx_path: str):
    md_text = Path(md_path).read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # 正文默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    i = 0
    in_code_block = False
    code_lang = ''
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # ---------- 代码块 ----------
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = stripped[3:].strip()
                code_buffer = []
            else:
                # 结束代码块
                in_code_block = False
                if code_lang.lower() == 'mermaid':
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run('[流程示意图：' + extract_mermaid_summary(code_buffer) + ']')
                    set_run_font(r, size=10, bold=True, color=(0x55, 0x55, 0x55))
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    for cl in code_buffer:
                        r = p.add_run(cl + '\n')
                        set_run_font(r, name='Consolas', size=9)
                code_buffer = []
                code_lang = ''
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ---------- 水平线 ----------
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'BFBFBF')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ---------- 标题 ----------
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            sizes = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}
            colors = {1: (0x1F, 0x3A, 0x68), 2: (0x1F, 0x3A, 0x68), 3: (0x2E, 0x5C, 0x9E), 4: (0x2E, 0x5C, 0x9E)}
            run = p.add_run(text)
            set_run_font(run, size=sizes.get(level, 11), bold=True, color=colors.get(level, (0, 0, 0)))
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # ---------- 表格 ----------
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            render_table(doc, table_lines)
            continue

        # ---------- 引用 ----------
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '18')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '1F3A68')
            pBdr.append(left)
            pPr.append(pBdr)
            add_runs_with_bold(p, text, base_size=11, base_color=(0x33, 0x33, 0x33), bold_color=(0x1F, 0x3A, 0x68))
            i += 1
            continue

        # ---------- 无序列表 ----------
        m_ul = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2
            text = m_ul.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.6 + indent * 0.6)
            add_runs_with_bold(p, text)
            i += 1
            continue

        # ---------- 有序列表 ----------
        m_ol = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if m_ol:
            text = m_ol.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs_with_bold(p, text)
            i += 1
            continue

        # ---------- 普通段落 ----------
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_runs_with_bold(p, stripped)
        else:
            # 空行：保留段落间距，不重复加段
            pass
        i += 1

    doc.save(docx_path)
    print(f'已生成 Word 文档：{docx_path}')


def extract_mermaid_summary(lines):
    """从 mermaid 节点中提取顺序串。"""
    nodes = []
    for ln in lines:
        for m in re.finditer(r'\[([^\]]+)\]', ln):
            nodes.append(m.group(1))
    seen = []
    for n in nodes:
        if n not in seen:
            seen.append(n)
    return ' → '.join(seen) if seen else '业务流程'


def render_table(doc, table_lines):
    """渲染 Markdown 表格。"""
    rows = []
    for ln in table_lines:
        # 去除首尾 |
        body = ln.strip()
        if body.startswith('|'):
            body = body[1:]
        if body.endswith('|'):
            body = body[:-1]
        cells = [c.strip() for c in body.split('|')]
        rows.append(cells)

    # 删除分隔行（| --- | --- |）
    if len(rows) >= 2 and all(re.match(r'^:?-+:?$', c) for c in rows[1]):
        header = rows[0]
        data = rows[2:]
    else:
        header = rows[0]
        data = rows[1:]

    col_count = max(len(header), max((len(r) for r in data), default=0))
    table = doc.add_table(rows=1 + len(data), cols=col_count)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, text in enumerate(header):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        add_runs_with_bold(p, text, base_size=10, base_color=(0xFF, 0xFF, 0xFF), bold_color=(0xFF, 0xFF, 0xFF))
        # 设置背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3A68')
        tcPr.append(shd)

    # 数据行
    for ri, row in enumerate(data, start=1):
        for ci in range(col_count):
            text = row[ci] if ci < len(row) else ''
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            add_runs_with_bold(p, text, base_size=10)


if __name__ == '__main__':
    src = r'g:\wqsPrint\snowy-admin-web\docs\条码打印系统解决方案.md'
    dst = r'g:\wqsPrint\snowy-admin-web\docs\条码打印系统解决方案.docx'
    if len(sys.argv) >= 3:
        src, dst = sys.argv[1], sys.argv[2]
    convert(src, dst)
